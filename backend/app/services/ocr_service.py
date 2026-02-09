"""
OCR service for extracting text from images and PDFs.
Config (engine, languages, gpu) is read from Settings table (type=ocr).
"""

import io
from pathlib import Path
from typing import List, Tuple
import tempfile

from PIL import Image
from sqlalchemy.orm import Session

from app.core.logging import get_logger
from app.core.settings_db import ensure_ocr_config
from app.models.assignment import SourceFormat

logger = get_logger()


class OCRService:
    """
    Service for extracting text from images and PDFs using EasyOCR.
    Config from Settings type=ocr (defaults in ensure_ocr_config).
    """

    def __init__(self, db: Session):
        self.db = db
        self._reader = None

    def _get_ocr_reader(self):
        """Get or create EasyOCR reader from DB config (Settings type=ocr)."""
        if self._reader is None:
            import easyocr
            rec = ensure_ocr_config(self.db)
            cfg = rec.config or {}
            languages = cfg.get("languages", ["en", "ch_sim"])
            gpu = cfg.get("gpu", False)
            self._reader = easyocr.Reader(languages, gpu=gpu)
            logger.info("Initialized EasyOCR with languages=%s, gpu=%s", languages, gpu)
        return self._reader

    def extract_text_from_image(self, image_path: str) -> str:
        """
        Extract text from an image file.

        Args:
            image_path: Path to the image file.

        Returns:
            Extracted text.
        """
        try:
            reader = self._get_ocr_reader()
            results = reader.readtext(image_path)

            # Extract text from results
            text_lines = [result[1] for result in results]
            extracted_text = "\n".join(text_lines)

            logger.info(f"Extracted {len(text_lines)} lines from image: {image_path}")
            return extracted_text

        except Exception as e:
            logger.error(f"OCR error for image {image_path}: {str(e)}")
            raise

    def extract_text_from_image_bytes(self, image_bytes: bytes) -> str:
        """
        Extract text from image bytes.

        Args:
            image_bytes: Raw image bytes.

        Returns:
            Extracted text.
        """
        try:
            # Save to temporary file for EasyOCR
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(image_bytes)
                tmp_path = tmp.name

            text = self.extract_text_from_image(tmp_path)

            # Clean up
            Path(tmp_path).unlink()

            return text

        except Exception as e:
            logger.error(f"OCR error for image bytes: {str(e)}")
            raise

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from a PDF file.

        First tries to extract embedded text, then falls back to OCR.

        Args:
            pdf_path: Path to the PDF file.

        Returns:
            Extracted text.
        """
        try:
            # Try to extract embedded text first
            import PyPDF2

            with open(pdf_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_parts = []

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                embedded_text = "\n\n".join(text_parts)

            # If we got meaningful text, return it
            if embedded_text and len(embedded_text.strip()) > 50:
                logger.info(f"Extracted embedded text from PDF: {pdf_path}")
                return embedded_text

            # Otherwise, fall back to OCR
            logger.info(f"No embedded text found, using OCR for PDF: {pdf_path}")
            return self._ocr_pdf(pdf_path)

        except Exception as e:
            logger.error(f"PDF extraction error for {pdf_path}: {str(e)}")
            # Try OCR as fallback
            return self._ocr_pdf(pdf_path)

    def _ocr_pdf(self, pdf_path: str) -> str:
        """
        Perform OCR on a PDF by converting to images.

        Args:
            pdf_path: Path to the PDF file.

        Returns:
            Extracted text.
        """
        try:
            from pdf2image import convert_from_path

            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=200)

            all_text = []
            reader = self._get_ocr_reader()

            for i, image in enumerate(images):
                # Convert PIL Image to bytes
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format="PNG")
                img_bytes = img_byte_arr.getvalue()

                # Save to temp file for EasyOCR
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(img_bytes)
                    tmp_path = tmp.name

                # OCR the page
                results = reader.readtext(tmp_path)
                page_text = "\n".join([r[1] for r in results])
                all_text.append(f"--- Page {i + 1} ---\n{page_text}")

                # Clean up
                Path(tmp_path).unlink()

            extracted_text = "\n\n".join(all_text)
            logger.info(f"OCR extracted text from {len(images)} PDF pages")
            return extracted_text

        except Exception as e:
            logger.error(f"PDF OCR error: {str(e)}")
            raise

    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from a Word document.

        Args:
            docx_path: Path to the DOCX file.

        Returns:
            Extracted text.
        """
        try:
            from docx import Document

            doc = Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)

            extracted_text = "\n".join(paragraphs)
            logger.info(f"Extracted text from DOCX: {docx_path}")
            return extracted_text

        except Exception as e:
            logger.error(f"DOCX extraction error for {docx_path}: {str(e)}")
            raise

    async def extract_text(
        self,
        file_path: str,
        source_format: SourceFormat,
    ) -> str:
        """
        Extract text from a file based on its format.

        Args:
            file_path: Path to the file.
            source_format: Format of the source file.

        Returns:
            Extracted text.
        """
        if source_format == SourceFormat.PDF:
            return self.extract_text_from_pdf(file_path)
        elif source_format in (SourceFormat.DOCX, SourceFormat.DOC):
            return self.extract_text_from_docx(file_path)
        elif source_format == SourceFormat.IMAGE:
            return self.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported format: {source_format}")


def get_ocr_service(db: Session) -> OCRService:
    """Get OCR service instance (config from Settings type=ocr)."""
    return OCRService(db)
