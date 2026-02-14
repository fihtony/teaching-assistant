"""
Convert markdown grading output to various formats (HTML, PDF, Word).

The markdown format uses special markup:
- ~~deleted text~~ → red strikethrough deletion
- {{added text}} → red addition
- ~~old~~{{new}} → red replacement (old struck, new underlined)

This converter transforms markdown to HTML and other formats,
applying color styling to corrections for display and export.
"""

import re
from typing import Dict, List, Tuple
from html import escape


class MarkdownGradingConverter:
    """Convert markdown grading output to various formats."""

    # Regex patterns for special markup
    REPLACEMENT_PATTERN = re.compile(r"~~([^~]+)~~\{\{([^}]+)\}\}")
    DELETION_PATTERN = re.compile(r"~~([^~]+)~~")
    ADDITION_PATTERN = re.compile(r"\{\{([^}]+)\}\}")

    @staticmethod
    def markdown_to_html(markdown_content: str, include_styling: bool = True) -> str:
        """
        Convert markdown grading output to HTML.

        Args:
            markdown_content: Markdown text with special correction markup
            include_styling: Whether to include inline CSS styles for colors

        Returns:
            HTML string with corrections styled in red
        """
        if not markdown_content:
            return ""

        # Preserve line breaks in "Dear X:" greeting format by marking them
        # This ensures that greetings like "Dear Student:\nWell done..." maintain line separation
        markdown_content_marked = markdown_content
        lines = markdown_content_marked.split("\n")
        preserved_lines = []
        for i, line in enumerate(lines):
            preserved_lines.append(line)
            # Mark line breaks after "Dear" greetings if followed by non-header content
            if line.strip().startswith("Dear") and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and not next_line.startswith("#"):
                    preserved_lines.append("<!--BREAK-->")
        markdown_content_marked = "\n".join(preserved_lines)

        html = markdown_content_marked

        # Convert special markup to HTML FIRST, before headers/formatting
        # Order matters: handle replacements first, then deletions and additions
        html = MarkdownGradingConverter._convert_replacements(html, include_styling)
        html = MarkdownGradingConverter._convert_deletions(html, include_styling)
        html = MarkdownGradingConverter._convert_additions(html, include_styling)

        # Apply inline formatting (bold, italic) BEFORE headers so that headers with bold/italic work
        # (e.g., ### **text** becomes ### <strong>text</strong> instead of <h3>**text**</h3>)
        html = MarkdownGradingConverter._apply_inline_formatting_to_all_lines(html)

        # Convert markdown headers - must happen before line-based formatting
        html = MarkdownGradingConverter._convert_markdown_headers(html)

        # Convert markdown formatting (paragraphs, lists)
        # Note: inline formatting already applied above, so _convert_markdown_formatting
        # should skip the _apply_inline_formatting step for already-formatted content
        html = MarkdownGradingConverter._convert_markdown_formatting(html)

        # Replace break markers with <br> tags
        html = html.replace("<!--BREAK-->", "<br>")

        return html

    @staticmethod
    def _convert_replacements(text: str, include_styling: bool = True) -> str:
        """Convert ~~old~~{{new}} to HTML with styling."""

        def replace_func(match):
            old_text = match.group(1)
            new_text = match.group(2)

            if include_styling:
                return (
                    f'<span style="color: #111; text-decoration: line-through; text-decoration-color: #dc2626;">{escape(old_text)}</span> '
                    f'<span style="color: #dc2626; text-decoration: underline;">{escape(new_text)}</span>'
                )
            else:
                return (
                    f'<span class="correction-deletion">{escape(old_text)}</span> '
                    f'<span class="correction-addition">{escape(new_text)}</span>'
                )

        return MarkdownGradingConverter.REPLACEMENT_PATTERN.sub(replace_func, text)

    @staticmethod
    def _convert_deletions(text: str, include_styling: bool = True) -> str:
        """Convert ~~deleted~~ to HTML with red strikethrough."""

        def replace_func(match):
            deleted_text = match.group(1)

            if include_styling:
                return f'<span style="color: #111; text-decoration: line-through; text-decoration-color: #dc2626;">{escape(deleted_text)}</span>'
            else:
                return (
                    f'<span class="correction-deletion">{escape(deleted_text)}</span>'
                )

        return MarkdownGradingConverter.DELETION_PATTERN.sub(replace_func, text)

    @staticmethod
    def _convert_additions(text: str, include_styling: bool = True) -> str:
        """Convert {{added}} to HTML with red text."""

        def replace_func(match):
            added_text = match.group(1)

            if include_styling:
                return f'<span style="color: red;">{escape(added_text)}</span>'
            else:
                return f'<span class="correction-addition">{escape(added_text)}</span>'

        return MarkdownGradingConverter.ADDITION_PATTERN.sub(replace_func, text)

    @staticmethod
    def _convert_markdown_headers(text: str) -> str:
        """Convert markdown headers (##, ###, etc.) to HTML headers."""
        # First, extract any headers that are incorrectly nested in list items (e.g., "- ### text")
        # and move them to their own line
        lines = text.split("\n")
        corrected_lines = []
        for line in lines:
            # Check if line starts with "- ###" or similar (header inside list) - extract header
            if re.match(r"^-\s+#{1,3}\s+\*\*.+?\*\*", line):
                # Extract the header part (e.g., "### **What You Did Well**")
                header_match = re.search(r"(#{1,3}\s+\*\*.+?\*\*)", line)
                if header_match:
                    header = header_match.group(1)
                    corrected_lines.append(header)
                else:
                    corrected_lines.append(line)
            else:
                corrected_lines.append(line)
        text = "\n".join(corrected_lines)

        # Convert ### to <h3>
        text = re.sub(r"^### (.+)$", r"<h3>\1</h3>", text, flags=re.MULTILINE)
        # Convert ## to <h2>
        text = re.sub(r"^## (.+)$", r"<h2>\1</h2>", text, flags=re.MULTILINE)
        # Convert # to <h1>
        text = re.sub(r"^# (.+)$", r"<h1>\1</h1>", text, flags=re.MULTILINE)
        return text

    @staticmethod
    def _convert_markdown_formatting(text: str) -> str:
        """
        Convert markdown formatting to HTML.
        Handles: lists, paragraphs, and spacing.
        NOTE: inline formatting (bold, italic, code) is already applied before this is called.
        Skips lines that are already HTML tags (e.g., headers).
        """
        lines = text.split("\n")
        html_lines = []
        in_list = False
        in_paragraph = False
        prev_was_empty = False
        current_list_item = None

        for i, line in enumerate(lines):
            stripped = line.rstrip()
            is_indented = line.startswith("  ") or line.startswith("\t")

            # Skip empty lines between list items and other content
            if not stripped:
                # Don't close list on empty lines - it might be between list items
                # Only emit empty paragraph break if not in list
                if in_paragraph and not in_list:
                    # Close continuation if we have it
                    if current_list_item:
                        current_list_item += "</li>"
                        html_lines.append(current_list_item)
                        current_list_item = None
                    html_lines.append("</p>")
                    in_paragraph = False
                    html_lines.append("")
                elif not in_list:
                    if current_list_item:
                        current_list_item += "</li>"
                        html_lines.append(current_list_item)
                        current_list_item = None
                    html_lines.append("")
                prev_was_empty = True
                continue

            # Check if this line is already an HTML tag (header, etc.) - skip formatting
            if stripped.startswith("<h") and (
                "</h1>" in stripped or "</h2>" in stripped or "</h3>" in stripped
            ):
                # Close any open tags first
                if current_list_item:
                    current_list_item += "</li>"
                    html_lines.append(current_list_item)
                    current_list_item = None
                if in_list:
                    html_lines.append("</ul>")
                    in_list = False
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False
                # Add the header as-is (inline formatting already applied)
                html_lines.append(stripped)
                prev_was_empty = False
                continue

            # Check if this is a bullet list item
            if stripped.startswith("- ") and not is_indented:
                # Close previous list item if exists
                if current_list_item:
                    current_list_item += "</li>"
                    html_lines.append(current_list_item)

                # Close paragraph BEFORE opening list (reorder: close p first, then open ul)
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False

                # Now open list if needed (only if not already open)
                if not in_list:
                    html_lines.append("<ul>")
                    in_list = True

                item_text = stripped[2:].strip()
                # Note: inline formatting already applied before headers were converted
                current_list_item = f"<li>{item_text}"
                prev_was_empty = False
                continue

            # Check if this is a continuation of the current list item (indented line)
            if is_indented and in_list and current_list_item:
                # This is a continuation of the list item - add to it with proper spacing
                continuation = stripped.strip()
                # Add double line break to maintain visual separation in HTML
                current_list_item += f"\n<br>\n{continuation}"
                prev_was_empty = False
                continue

            # Close list if we're transitioning to non-list content
            if in_list:
                if current_list_item:
                    current_list_item += "</li>"
                    html_lines.append(current_list_item)
                    current_list_item = None
                html_lines.append("</ul>")
                in_list = False

            # Regular paragraph text
            # Note: inline formatting already applied before headers
            if not in_paragraph:
                html_lines.append(f"<p>{stripped}")
                in_paragraph = True
            else:
                html_lines.append(f"{stripped}")
            prev_was_empty = False

        # Close any open tags
        if current_list_item:
            current_list_item += "</li>"
            html_lines.append(current_list_item)
        if in_list:
            html_lines.append("</ul>")
        if in_paragraph:
            html_lines.append("</p>")

        result = "\n".join(html_lines)
        return result

    @staticmethod
    def _apply_inline_formatting_to_all_lines(text: str) -> str:
        """
        Apply inline formatting (bold, italic, code) to all non-HTML lines.
        This is called before header conversion so that headers with bold/italic work.
        """
        lines = text.split("\n")
        result_lines = []

        for line in lines:
            # Skip lines that are already HTML tags (are pure HTML)
            if line.strip().startswith("<") and line.strip().endswith(">"):
                result_lines.append(line)
            else:
                # Apply inline formatting to this line
                formatted_line = MarkdownGradingConverter._apply_inline_formatting(line)
                result_lines.append(formatted_line)

        return "\n".join(result_lines)

    @staticmethod
    def _apply_inline_formatting(text: str) -> str:
        """Apply inline formatting (bold, italic, code) to text."""
        # Convert **bold** to <strong>
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        # Convert *italic* to <em> (but not **bold**)
        text = re.sub(r"(?<!\*)\*(.+?)\*(?!\*)", r"<em>\1</em>", text)
        # Convert `code` to <code>
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        return text

    @staticmethod
    def get_css_styles() -> str:
        """
        Return CSS styles for correction markup when include_styling=False.
        Use this with HTML output when you want external CSS instead of inline styles.
        """
        return """
        <style>
            .correction-deletion {
                color: red;
                text-decoration: line-through;
            }
            .correction-addition {
                color: red;
            }
            .original-text {
                color: black;
            }
            .explanation {
                color: black;
                font-style: italic;
                margin-left: 1em;
            }
        </style>
        """

    @staticmethod
    def markdown_to_docx(markdown_content: str) -> bytes:
        """
        Convert markdown grading output to DOCX format with tracking mode enabled.

        Converts markdown markup to Word formatting:
        - ~~deleted~~ → Red text with strikethrough (tracked deletion)
        - {{added}} → Red text (tracked insertion)
        - ~~old~~{{new}} → Red replacement (old struck, new underlined)

        Args:
            markdown_content: Markdown text with special correction markup

        Returns:
            DOCX file as bytes with tracking changes enabled
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )

        doc = Document()

        # Enable Track Changes in the document
        MarkdownGradingConverter._enable_track_changes(doc)

        # Parse and add content
        lines = markdown_content.split("\n")
        for line in lines:
            if line.startswith("### "):
                # Heading level 3 with formatting
                heading_text = line[4:].strip()
                MarkdownGradingConverter._add_formatted_paragraph(
                    doc, heading_text, style="Heading 3"
                )
            elif line.startswith("## "):
                # Heading level 2 with formatting
                heading_text = line[3:].strip()
                MarkdownGradingConverter._add_formatted_paragraph(
                    doc, heading_text, style="Heading 2"
                )
            elif line.startswith("# "):
                # Heading level 1 with formatting
                heading_text = line[2:].strip()
                MarkdownGradingConverter._add_formatted_paragraph(
                    doc, heading_text, style="Heading 1"
                )
            elif line.startswith("- "):
                # Bullet point with formatting
                bullet_text = line[2:].strip()
                MarkdownGradingConverter._add_formatted_paragraph(
                    doc, bullet_text, style="List Bullet"
                )
            elif line.strip():
                # Regular paragraph
                MarkdownGradingConverter._add_formatted_paragraph(doc, line.strip())

        # Convert document to bytes
        import io

        byte_stream = io.BytesIO()
        doc.save(byte_stream)
        return byte_stream.getvalue()

    @staticmethod
    def _enable_track_changes(doc):
        """
        Enable Track Changes in the Word document.

        This allows Word to display changes as tracked revisions when opened.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Access document settings
        settings = doc.settings
        settings_element = settings._element

        # Add revision tracking element if not present
        track_changes = OxmlElement("w:trackRevisions")
        settings_element.append(track_changes)

    @staticmethod
    def _add_formatted_paragraph(doc, text: str, style: str = None):
        """
        Add a paragraph to the document with correction markup converted to formatting.

        Handles ~~deletion~~, {{addition}}, and ~~old~~{{new}} markup.
        Applies red text color and strikethrough/underline for tracked changes visibility.

        Args:
            doc: python-docx Document object
            text: Text content with markup
            style: Optional paragraph style (e.g., 'Heading 1', 'List Bullet')
        """
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Add paragraph with optional style
        if style:
            p = doc.add_paragraph(style=style)
        else:
            p = doc.add_paragraph()

        # Process the text and apply formatting
        remaining = text

        while remaining:
            # Check for replacement pattern first: ~~old~~{{new}}
            match = MarkdownGradingConverter.REPLACEMENT_PATTERN.search(remaining)

            if match:
                # Add text before the match
                if match.start() > 0:
                    run = p.add_run(remaining[: match.start()])
                    run.font.color.rgb = RGBColor(17, 17, 17)  # Dark gray/black

                # Add deleted (old) text in red with strikethrough
                deleted_run = p.add_run(match.group(1))
                # Set red color (via XML) and strikethrough (via both API and XML)
                deleted_run.font.strikethrough = True  # API call for compatibility
                MarkdownGradingConverter._apply_red_formatting(
                    deleted_run, strikethrough=True
                )

                # Add added (new) text in red with underline
                added_run = p.add_run(match.group(2))
                # Set red color (via XML) and underline (via API)
                added_run.underline = True
                MarkdownGradingConverter._apply_red_formatting(
                    added_run, strikethrough=False
                )

                remaining = remaining[match.end() :]
            else:
                # Check for simple deletion: ~~text~~
                match = MarkdownGradingConverter.DELETION_PATTERN.search(remaining)

                if match:
                    # Add text before the match
                    if match.start() > 0:
                        run = p.add_run(remaining[: match.start()])
                        run.font.color.rgb = RGBColor(17, 17, 17)  # Dark gray/black

                    # Add deleted text in red with strikethrough
                    deleted_run = p.add_run(match.group(1))
                    # Set red color (via XML) and strikethrough (via both API and XML)
                    deleted_run.font.strikethrough = True  # API call for compatibility
                    MarkdownGradingConverter._apply_red_formatting(
                        deleted_run, strikethrough=True
                    )

                    remaining = remaining[match.end() :]
                else:
                    # Check for simple addition: {{text}}
                    match = MarkdownGradingConverter.ADDITION_PATTERN.search(remaining)

                    if match:
                        # Add text before the match
                        if match.start() > 0:
                            run = p.add_run(remaining[: match.start()])
                            run.font.color.rgb = RGBColor(17, 17, 17)  # Dark gray/black

                        # Add added text in red
                        added_run = p.add_run(match.group(1))
                        # Set red color using XML manipulation
                        MarkdownGradingConverter._apply_red_formatting(
                            added_run, strikethrough=False
                        )
                        added_run.underline = True

                        remaining = remaining[match.end() :]
                    else:
                        # No more markup, add remaining text
                        run = p.add_run(remaining)
                        run.font.color.rgb = RGBColor(17, 17, 17)  # Dark gray/black
                        remaining = ""

    @staticmethod
    def _apply_red_formatting(run, strikethrough: bool = False):
        """
        Apply red formatting to a run using direct XML manipulation.

        This function sets:
        - Text color to red (C00000 - Word standard red)
        - Strikethrough if requested

        Uses pure XML manipulation to ensure formatting is preserved in Word export.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        rPr = run._element.get_or_add_rPr()

        # Step 1: Remove any existing color elements to avoid conflicts
        for color_elem in list(rPr.findall(qn("w:color"))):
            rPr.remove(color_elem)

        # Step 2: Remove any existing strikethrough elements to avoid conflicts
        for strike_elem in list(rPr.findall(qn("w:strike"))):
            rPr.remove(strike_elem)

        # Step 3: Add red color element
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), "C00000")  # Word standard red
        rPr.append(color_elem)

        # Step 4: Add strikethrough element if requested
        if strikethrough:
            strike_elem = OxmlElement("w:strike")
            rPr.append(strike_elem)

    @staticmethod
    def extract_sections(markdown_content: str) -> Dict[str, str]:
        """
        Extract main sections from markdown grading output.

        Returns:
            Dictionary with keys: 'revised_essay', 'detailed_corrections', 'teacher_comments'
        """
        sections = {
            "revised_essay": "",
            "detailed_corrections": "",
            "teacher_comments": "",
        }

        lines = markdown_content.split("\n")
        current_section = None
        current_content = []

        for line in lines:
            if line.startswith("## Revised Essay"):
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = "revised_essay"
                current_content = []
            elif line.startswith("## Detailed Corrections"):
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = "detailed_corrections"
                current_content = []
            elif line.startswith("## Teacher's Comments"):
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = "teacher_comments"
                current_content = []
            elif current_section:
                current_content.append(line)

        # Add final section
        if current_section and current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections
