"""
Export service for generating graded documents.
"""

import io
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple
import json

from bs4 import BeautifulSoup
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfgen import canvas

from app.core.config import get_storage_path
from app.core.logging import get_logger
from app.models import Assignment, SourceFormat
from app.schemas.assignment import GradingResult, ExportFormat

logger = get_logger()


class ExportService:
    """
    Service for exporting graded assignments.

    Generates PDF or DOCX documents with inline annotations.
    """

    # Annotation colors
    CORRECT_COLOR = colors.green
    INCORRECT_COLOR = colors.red
    COMMENT_COLOR = colors.blue
    ENCOURAGEMENT_COLOR = colors.purple

    def __init__(self):
        self.graded_dir = get_storage_path("graded")

    def determine_export_format(
        self,
        source_format: SourceFormat,
        requested_format: ExportFormat,
    ) -> SourceFormat:
        """
        Determine the actual export format.

        Args:
            source_format: Original file format.
            requested_format: Requested export format.

        Returns:
            Actual format to use.
        """
        if requested_format == ExportFormat.AUTO:
            # Return same format as source, with fallbacks
            if source_format in (SourceFormat.PDF, SourceFormat.IMAGE):
                return SourceFormat.PDF
            elif source_format in (SourceFormat.DOCX, SourceFormat.DOC):
                return SourceFormat.DOCX
            elif source_format == SourceFormat.TXT:
                return SourceFormat.DOCX  # Export graded TXT as DOCX
            else:
                return SourceFormat.PDF
        elif requested_format == ExportFormat.PDF:
            return SourceFormat.PDF
        elif requested_format == ExportFormat.DOCX:
            return SourceFormat.DOCX
        else:
            return SourceFormat.PDF

    def export_to_pdf(
        self,
        assignment: Assignment,
        grading_result: GradingResult,
    ) -> Tuple[bytes, str]:
        """
        Export graded assignment to PDF.

        Args:
            assignment: The assignment.
            grading_result: Grading results.

        Returns:
            Tuple of (PDF bytes, filename).
        """
        buffer = io.BytesIO()

        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        # Styles
        styles = getSampleStyleSheet()
        title_style = styles["Heading1"]
        heading_style = styles["Heading2"]
        normal_style = styles["Normal"]

        # Custom styles for annotations
        correct_style = ParagraphStyle(
            "Correct",
            parent=normal_style,
            textColor=colors.green,
            fontSize=10,
        )
        incorrect_style = ParagraphStyle(
            "Incorrect",
            parent=normal_style,
            textColor=colors.red,
            fontSize=10,
        )
        comment_style = ParagraphStyle(
            "Comment",
            parent=normal_style,
            textColor=colors.blue,
            fontSize=9,
            leftIndent=20,
        )
        encouragement_style = ParagraphStyle(
            "Encouragement",
            parent=normal_style,
            textColor=colors.purple,
            fontSize=12,
            fontName="Helvetica-Bold",
        )

        # Build document content
        content = []

        # Title
        title = assignment.student_name or assignment.original_filename
        content.append(Paragraph(f"Graded Assignment: {title}", title_style))
        content.append(Spacer(1, 12))

        # Metadata
        content.append(
            Paragraph(
                f"Graded on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style
            )
        )
        content.append(Spacer(1, 24))

        # Original text with inline annotations
        content.append(Paragraph("Student's Work with Annotations:", heading_style))
        content.append(Spacer(1, 12))

        # Process each grading item
        for item in grading_result.items:
            # Question header
            content.append(
                Paragraph(
                    f"<b>Question {item.question_number}</b> ({item.question_type.value})",
                    normal_style,
                )
            )

            # Student answer
            answer_text = f"Student's Answer: {item.student_answer}"
            if item.is_correct:
                content.append(Paragraph(f"✓ {answer_text}", correct_style))
            else:
                content.append(Paragraph(f"✗ {answer_text}", incorrect_style))
                if item.correct_answer:
                    content.append(
                        Paragraph(
                            f"Correct Answer: {item.correct_answer}", correct_style
                        )
                    )

            # Comment
            if item.comment:
                content.append(Paragraph(f"💬 {item.comment}", comment_style))

            content.append(Spacer(1, 12))

        # Section scores with encouragements
        content.append(Spacer(1, 24))
        content.append(Paragraph("Section Scores:", heading_style))
        content.append(Spacer(1, 12))

        for section_name, scores in grading_result.section_scores.items():
            score_text = f"{section_name.upper()}: {scores.correct}/{scores.total}"
            content.append(Paragraph(score_text, normal_style))

            if scores.encouragement:
                content.append(
                    Paragraph(f"🌟 {scores.encouragement}", encouragement_style)
                )

            content.append(Spacer(1, 6))

        # Overall comment
        if grading_result.overall_comment:
            content.append(Spacer(1, 24))
            content.append(Paragraph("Overall Feedback:", heading_style))
            content.append(Paragraph(grading_result.overall_comment, normal_style))

        # Build PDF
        doc.build(content)

        # Get bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()

        # Generate filename
        base_name = Path(assignment.stored_filename).stem
        filename = f"{base_name}_graded.pdf"

        # Save to graded directory
        output_path = self.graded_dir / filename
        with open(output_path, "wb") as f:
            f.write(pdf_bytes)

        logger.info(f"Exported graded PDF: {filename}")

        return pdf_bytes, filename

    def export_to_docx(
        self,
        assignment: Assignment,
        grading_result: GradingResult,
    ) -> Tuple[bytes, str]:
        """
        Export graded assignment to DOCX.

        Args:
            assignment: The assignment.
            grading_result: Grading results.

        Returns:
            Tuple of (DOCX bytes, filename).
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_COLOR_INDEX

        doc = Document()

        # Title
        title = assignment.student_name or assignment.original_filename
        doc.add_heading(f"Graded Assignment: {title}", 0)

        # Metadata
        doc.add_paragraph(f"Graded on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # Student's work with annotations
        doc.add_heading("Student's Work with Annotations:", level=1)

        for item in grading_result.items:
            # Question header
            p = doc.add_paragraph()
            run = p.add_run(
                f"Question {item.question_number} ({item.question_type.value})"
            )
            run.bold = True

            # Student answer with color
            p = doc.add_paragraph()
            if item.is_correct:
                run = p.add_run(f"✓ Student's Answer: {item.student_answer}")
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                run = p.add_run(f"✗ Student's Answer: {item.student_answer}")
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red

                if item.correct_answer:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Correct Answer: {item.correct_answer}")
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green

            # Comment in blue
            if item.comment:
                p = doc.add_paragraph()
                run = p.add_run(f"💬 {item.comment}")
                run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                run.font.size = Pt(9)

            doc.add_paragraph()

        # Section scores
        doc.add_heading("Section Scores:", level=1)

        for section_name, scores in grading_result.section_scores.items():
            p = doc.add_paragraph()
            p.add_run(f"{section_name.upper()}: {scores.correct}/{scores.total}")

            if scores.encouragement:
                p = doc.add_paragraph()
                run = p.add_run(f"🌟 {scores.encouragement}")
                run.font.color.rgb = RGBColor(128, 0, 128)  # Purple
                run.bold = True

        # Overall comment
        if grading_result.overall_comment:
            doc.add_heading("Overall Feedback:", level=1)
            doc.add_paragraph(grading_result.overall_comment)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        # Generate filename
        base_name = Path(assignment.stored_filename).stem
        filename = f"{base_name}_graded.docx"

        # Save to graded directory
        output_path = self.graded_dir / filename
        with open(output_path, "wb") as f:
            f.write(docx_bytes)

        logger.info(f"Exported graded DOCX: {filename}")

        return docx_bytes, filename

    def _html_to_pdf(self, html_content: str) -> bytes:
        """Convert HTML fragment to PDF bytes (for HTML-only grading result)."""
        try:
            from xhtml2pdf import pisa
        except ImportError:
            logger.warning("xhtml2pdf not installed; cannot export HTML as PDF")
            raise ValueError(
                "PDF export from HTML requires xhtml2pdf. Install with: pip install xhtml2pdf"
            )

        # Wrap in full document; use only simple CSS (xhtml2pdf does not support attribute selectors)
        doc = f"""<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><meta charset="utf-8"/><style>
del {{ text-decoration: line-through; color: #dc2626; }}
.correction {{ color: #dc2626; }}
h2, h3 {{ font-weight: bold; margin-top: 1em; }}
</style></head>
<body>{html_content}</body>
</html>"""
        output = io.BytesIO()
        pisa_status = pisa.CreatePDF(doc, dest=output, encoding="utf-8")
        if pisa_status.err:
            logger.error("xhtml2pdf reported errors during conversion")
        pdf_bytes = output.getvalue()
        output.close()
        return pdf_bytes

    def _html_to_docx(self, html_content: str, title: str) -> bytes:
        """Convert HTML fragment to DOCX bytes while preserving formatting from inline styles.

        Extracts colors and strikethrough from HTML span styles and applies them to Word formatting.
        """
        from docx import Document
        from docx.shared import RGBColor
        import re

        doc = Document()
        doc.add_heading(f"Graded Assignment: {title}", 0)
        doc.add_paragraph(f"Graded on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        soup = BeautifulSoup(html_content, "html.parser")

        for el in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol"]):
            if el.name in ("h1", "h2", "h3"):
                level = {"h1": 0, "h2": 1, "h3": 2}[el.name]
                # Process heading with styles
                p = doc.add_paragraph(style=f"Heading {level + 1}")
                self._process_element_with_styles(p, el)
            elif el.name == "p":
                # Process paragraph with styles
                p = doc.add_paragraph()
                self._process_element_with_styles(p, el)
            elif el.name in ("ul", "ol"):
                for li in el.find_all("li", recursive=False):
                    # Process list item with styles
                    p = doc.add_paragraph(style="List Bullet")
                    self._process_element_with_styles(p, li)

        # If no block elements found, add whole body as paragraphs
        if len(doc.paragraphs) <= 2:
            body = soup.get_text(separator="\n", strip=True)
            if body:
                for line in body.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    def _process_element_with_styles(self, paragraph, element):
        """Process an HTML element and apply styles (colors, strikethrough) to Word paragraph.

        Recursively processes child elements and extracts style information from spans.
        Handles text-decoration-color for red strikethrough lines.
        """
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import re

        for content in element.children:
            if isinstance(content, str):
                # Plain text
                text = str(content).strip()
                if text:
                    run = paragraph.add_run(text)
            else:
                # HTML element
                tag_name = content.name

                if tag_name == "span":
                    # Extract style attribute
                    style = content.get("style", "")
                    text = content.get_text()

                    if text:
                        run = paragraph.add_run(text)

                        # Check if this is strikethrough with red decoration (deletion in Word)
                        has_strikethrough = "text-decoration: line-through" in style
                        has_red_strike_color = "text-decoration-color: #dc2626" in style

                        # For Word export: use light grey + italic for deleted text (strikethrough with red color indicator)
                        # This is because in Word, strikethrough color automatically follows text color
                        if has_strikethrough and has_red_strike_color:
                            # Use light grey for text and add italic for emphasis
                            run.font.color.rgb = RGBColor(160, 160, 160)  # Light grey
                            run.italic = True
                        else:
                            # Extract and apply text color normally for non-deleted text
                            color_match = re.search(
                                r"color:\s*#([0-9A-Fa-f]{6})", style
                            )
                            if color_match:
                                hex_color = color_match.group(1)
                                r = int(hex_color[0:2], 16)
                                g = int(hex_color[2:4], 16)
                                b = int(hex_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)

                        # Apply strikethrough if present
                        if "text-decoration: line-through" in style:
                            run.font.strikethrough = True
                            # Also apply at XML level to ensure it shows in Word
                            rPr = run._element.rPr
                            if rPr is None:
                                rPr = OxmlElement("w:rPr")
                                run._element.insert(0, rPr)
                            # Create strike element
                            strike = OxmlElement("w:strike")
                            # Remove existing strike if present
                            for existing_strike in rPr.findall(qn("w:strike")):
                                rPr.remove(existing_strike)
                            rPr.append(strike)

                        # Apply underline if present
                        if "text-decoration: underline" in style:
                            run.underline = True

                        # Apply bold if needed
                        if "font-weight: bold" in style:
                            run.bold = True

                elif tag_name == "em":
                    text = content.get_text()
                    if text:
                        run = paragraph.add_run(text)
                        run.italic = True

                elif tag_name == "br":
                    paragraph.add_run("\n")

                else:
                    # Recursively process other elements
                    text = content.get_text()
                    if text:
                        run = paragraph.add_run(text)

    async def export(
        self,
        assignment: Assignment,
        grading_result: GradingResult,
        export_format: ExportFormat = ExportFormat.AUTO,
    ) -> Tuple[bytes, str, str]:
        """
        Export a graded assignment.

        Args:
            assignment: The assignment.
            grading_result: Grading results.
            export_format: Desired export format.

        Returns:
            Tuple of (file bytes, filename, content_type).
        """
        # When result is HTML-only (no structured items), convert to requested format
        if grading_result.html_content and not grading_result.items:
            base_name = Path(assignment.stored_filename).stem
            title = assignment.student_name or assignment.original_filename
            actual_format = self.determine_export_format(
                assignment.source_format,
                export_format,
            )
            if actual_format == SourceFormat.PDF:
                content = self._html_to_pdf(grading_result.html_content)
                filename = f"{base_name}_graded.pdf"
                content_type = "application/pdf"
            else:
                content = self._html_to_docx(
                    grading_result.html_content, title or base_name
                )
                filename = f"{base_name}_graded.docx"
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return content, filename, content_type

        actual_format = self.determine_export_format(
            assignment.source_format,
            export_format,
        )

        if actual_format == SourceFormat.PDF:
            content, filename = self.export_to_pdf(assignment, grading_result)
            content_type = "application/pdf"
        else:
            content, filename = self.export_to_docx(assignment, grading_result)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Update assignment with graded filename
        # graded_filename is stored on ai_grading record in API

        return content, filename, content_type


# Global export service instance
_export_service: Optional[ExportService] = None


def get_export_service() -> ExportService:
    """Get the global export service instance."""
    global _export_service
    if _export_service is None:
        _export_service = ExportService()
    return _export_service
